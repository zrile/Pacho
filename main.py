import docx
doc = docx.Document('probno1.docx')
print (len(doc.paragraphs))
for i in range (len(doc.paragraphs)):
    print (doc.paragraphs[i].text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print (cell.text)